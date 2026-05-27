from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DIAGRAMS = DOCS / "diagrams"
OUTPUT = DOCS / "AI_ASSISTED_TEST_AUTOMATION_PROJECT_REPORT_FIXED_TABLES.docx"

BLUE = RGBColor(23, 54, 93)
GOLD = RGBColor(156, 101, 0)
DARK = RGBColor(45, 45, 45)
GRAY = RGBColor(90, 90, 90)


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    if level == 1:
        set_font(r, 18, True, BLUE)
    elif level == 2:
        set_font(r, 14, True, GOLD)
    else:
        set_font(r, 12, True, DARK)
    return p


def add_para(doc, text, bold=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    set_font(r, 11, bold, color)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        set_font(r, 11)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        set_font(r, 11)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        from docx.oxml import OxmlElement

        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill", fill)


def style_cell(cell, header=False, fill=None, font_size=9):
    if fill:
        shade_cell(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            set_font(run, font_size, header, BLUE if header else None)


def add_table(doc, headers, rows, fill="D9EAF7", font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        style_cell(cell, True, fill, font_size + 1)

    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for col_index, value in enumerate(row):
            cells[col_index].text = str(value)
            style_cell(cells[col_index], False, "F7F7F7" if row_index % 2 else None, font_size)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_diagram(doc, file_name, caption, width=7.4):
    image = DIAGRAMS / file_name
    if not image.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run()
    run.add_picture(str(image), width=Inches(width))

    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(6)
    r = c.add_run(caption)
    set_font(r, 9, True, GRAY)


def build():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.text = ""

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    title_run = title.add_run("Demo Web Shop Automation Framework")
    set_font(title_run, 22, True, BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle_run = subtitle.add_run("AI-Assisted Test Automation Project Report for Customer Order Placement User Story")
    set_font(subtitle_run, 13, True, GOLD)

    add_heading(doc, "Introduction", 1)
    add_para(
        doc,
        "This document presents the complete end-to-end automation work completed for the Demo Web Shop project. "
        "The automation framework was enhanced for the customer order placement user story: as a customer, the user "
        "should be able to log in, view products, add a product to the cart, complete checkout, and successfully place an order.",
    )
    add_para(
        doc,
        "AI was used as a testing accelerator to generate scenarios, identify edge cases, prioritize risk-based coverage, "
        "improve automation stability, optimize regression execution, and prepare structured documentation. Human validation "
        "was still used to review failures, verify real application behavior, and approve the final result.",
    )

    add_table(
        doc,
        ["Metric", "Final Result"],
        [
            ["Total tests executed", "31"],
            ["Passed", "31"],
            ["Failed / Errors / Skipped", "0 / 0 / 0"],
            ["Final execution time", "681.1 seconds"],
            ["Execution command", 'mvn -q test "-Dheadless=true"'],
        ],
        "D9EAD3",
        10,
    )
    add_diagram(doc, "ai_testing_process.png", "Figure 1: AI-assisted testing process with human review and final execution evidence.", 8.2)

    add_heading(doc, "User Story And Scope", 1)
    add_para(
        doc,
        "User Story: As a customer, I want to log in, view products, add a product to the cart, and complete checkout, so that I can successfully place an order.",
        True,
        DARK,
    )
    add_bullets(
        doc,
        [
            "Validate the complete customer journey from login to order confirmation.",
            "Identify functional, UI, edge, and non-functional test cases.",
            "Keep Maven, Selenium, TestNG, and the hybrid framework structure unchanged.",
            "Refactor page and test classes to support frequent UI changes.",
            "Add smart UI handling, fallback locators, explicit waits, Log4j, Extent Reports, screenshots on failure, and Excel data-driven testing.",
            "Prioritize critical tests based on business risk and reduce unnecessary regression execution time.",
        ],
    )
    add_diagram(doc, "customer_order_flow.png", "Figure 2: Customer order placement journey covered by the automation suite.", 7.8)

    add_heading(doc, "Framework And Tool Selection Rationale", 1)
    add_diagram(doc, "framework_architecture.png", "Figure 3: Hybrid Maven, Selenium, TestNG, POM, Excel, logging, and reporting structure.", 7.8)
    add_table(
        doc,
        ["Tool / Library", "Purpose", "Why It Was Selected"],
        [
            ["Maven", "Build and dependency management", "Standard lifecycle, Surefire integration, and easy suite execution."],
            ["Selenium WebDriver", "Browser automation", "Best fit for forms, buttons, links, checkout pages, alerts, and dynamic UI."],
            ["TestNG", "Test framework", "Supports suites, groups, priority, data providers, setup/teardown, and reporting."],
            ["Apache POI", "Excel data handling", "Business-friendly data-driven testing using TestData.xlsx."],
            ["Log4j 2", "Execution logging", "Debugging and audit trail during execution."],
            ["Extent Reports", "HTML reporting", "Readable reports with steps, screenshots, pass/fail status, and environment details."],
            ["Page Object Model", "Maintainability", "Separates page logic from test logic and reduces duplicated locator handling."],
        ],
        "D9EAF7",
        8,
    )

    add_heading(doc, "Process Followed", 1)
    add_numbers(
        doc,
        [
            "Reviewed the existing Maven project structure, utilities, page classes, test classes, and TestNG XML files.",
            "Identified unstable areas: hard-coded credentials, brittle locators, dynamic checkout sections, fixed product assumptions, and missing alert handling.",
            "Refactored base page and page objects for smart waits, fallback locators, JavaScript click fallback, scrolling, and optional fields.",
            "Added reusable business flows for registration, login, product selection, cart, and checkout.",
            "Generated Excel workbook support for test data and updated data providers.",
            "Added AI-generated functional, UI, edge, non-functional, and risk-based tests.",
            "Separated critical, user-story, and full regression suites.",
            "Ran the full suite, fixed failures, and verified the final result: 31 tests passed and 0 failed.",
        ],
    )

    add_heading(doc, "Outputs Generated By AI", 1)
    add_table(
        doc,
        ["Output", "Description"],
        [
            ["User story test plan", "Functional, UI, edge, non-functional, and risk-based coverage."],
            ["Critical TestNG suite", "Small high-impact suite for fast validation."],
            ["User-story TestNG suite", "Focused suite for user-story level coverage."],
            ["Refactored page objects", "Smart locators, fallback locators, and dynamic checkout handling."],
            ["Data-driven Excel workbook", "Auto-generated TestData.xlsx with required sheets."],
            ["Test data strategy", "Checkout, login, registration, and product search data."],
            ["Optimized regression approach", "Critical-first execution and stable reusable flows."],
            ["Bug observations", "Application behavior and automation-handled issues."],
            ["Final green evidence", "31 tests, 0 failures, 0 skipped."],
        ],
        "D9EAD3",
        8,
    )

    add_heading(doc, "AI-Generated Test Scenarios And Test Cases", 1)
    add_para(
        doc,
        "This section is retained as evidence for the submission. It shows how the user story was converted into functional, UI, edge, non-functional, and regression coverage.",
    )
    add_heading(doc, "Scenario Coverage Matrix", 2)
    add_table(
        doc,
        ["ID", "Scenario", "Type", "Priority", "Risk", "Expected Result"],
        [
            ["TC-001", "Register or prepare customer account", "Functional", "High", "Login requires valid account", "Customer account is available."],
            ["TC-002", "Login with valid credentials", "Functional", "Critical", "Customer cannot purchase without login", "User logs in and logout link is visible."],
            ["TC-003", "Login with invalid email", "Edge", "High", "Authentication validation", "Error message is shown."],
            ["TC-004", "Login with invalid password", "Edge", "High", "Account security validation", "Error message is shown."],
            ["TC-005", "Search for product by keyword", "Functional", "Critical", "Product discovery drives purchase", "Search results are displayed."],
            ["TC-006", "Validate product list UI", "UI", "High", "Product card affects decision", "Product names and prices are visible."],
            ["TC-007", "View product details", "Functional/UI", "Critical", "Customer must inspect product", "Title, price, and Add to Cart are visible."],
            ["TC-008", "Add product to cart", "Functional", "Critical", "Revenue path breaks if cart fails", "Product is present in cart."],
            ["TC-009", "Validate cart summary", "UI", "High", "Customer must review cart", "Rows, quantity, subtotal, terms, and checkout controls are visible."],
            ["TC-010", "Proceed to checkout", "Functional", "Critical", "Checkout entry is business critical", "Checkout page loads."],
            ["TC-011", "Enter billing address", "Functional", "Critical", "Order cannot be placed without billing", "Billing data is accepted."],
            ["TC-012", "Select shipping method", "Functional", "High", "Shipping required for order", "Shipping method is selected."],
            ["TC-013", "Select payment method", "Functional", "High", "Payment required for order", "Payment method is selected."],
            ["TC-014", "Confirm order", "Functional", "Critical", "Final order placement point", "Confirmation page is displayed."],
            ["TC-015", "Missing billing details", "Edge", "High", "Required-field validation", "Validation error is displayed."],
            ["TC-016", "Empty product search", "Edge", "Medium", "Search validation behavior", "Alert or validation message is handled."],
            ["TC-017", "Product discovery response time", "Non-functional", "Medium", "Slow search affects experience", "Search completes within threshold."],
            ["TC-018", "Multiple product cart flow", "Regression", "Medium", "Cart should support repeat add flows", "Cart remains usable and checkout succeeds."],
        ],
        "FFF2CC",
        7,
    )

    add_heading(doc, "Detailed Test Cases", 2)
    add_table(
        doc,
        ["Test Case", "Precondition", "Test Steps", "Expected Result"],
        [
            ["Valid login", "Registered user exists", "Open login page; enter email and password; click Login.", "User is logged in and logout link appears."],
            ["Invalid login", "Application is available", "Open login page; enter invalid credentials; click Login.", "Authentication error appears."],
            ["Product search", "User is on home page", "Enter product keyword; click Search.", "Matching products are displayed."],
            ["Product details UI", "Search results are available", "Open a product from the result list.", "Product title, price, and Add to Cart are visible."],
            ["Add to cart", "Product details page is open", "Click Add to Cart; open shopping cart.", "Selected product appears in the cart."],
            ["Cart UI validation", "Cart has a product", "Inspect cart row, quantity field, terms checkbox, and checkout button.", "All critical cart controls are visible."],
            ["Checkout happy path", "Logged-in user has product in cart", "Accept terms; click Checkout; complete billing, shipping, payment, and confirmation.", "Order confirmation page is displayed."],
            ["Missing billing edge case", "Checkout page is open", "Leave required billing details empty and continue.", "Validation message appears."],
            ["Empty search edge case", "User is on home page", "Keep search box empty and click Search.", "Browser alert is handled and validated."],
            ["Search performance", "Application is available", "Run product search and measure response time.", "Response stays within configured threshold."],
        ],
        "E4DFEC",
        8,
    )

    add_heading(doc, "Optimized Test Cases With AI Reasoning", 2)
    add_table(
        doc,
        ["Optimized Test", "AI Reasoning"],
        [
            ["Critical order placement path", "Highest business value because it validates login, product discovery, cart, checkout, and confirmation in one journey."],
            ["UI element validation at each step", "UI changes frequently, so validating key controls detects breaking changes early."],
            ["Missing billing details edge test", "Required-field validation is important but should not block the critical smoke path."],
            ["Product search performance check", "Lightweight non-functional validation gives early signal of response degradation."],
            ["Login negative tests", "Authentication failures are high risk but should remain separate from purchase happy path."],
            ["Empty search alert handling", "Prevents browser alert from breaking execution and documents current application behavior."],
            ["Cart update and removal tests", "Medium-risk regression checks kept outside the smallest critical suite."],
            ["Full E2E tests", "Valuable for regression but slower, so not included in the smallest critical suite."],
        ],
        "D9EAD3",
        8,
    )

    add_heading(doc, "Functional, UI, Edge, And Non-Functional Coverage", 1)
    add_table(
        doc,
        ["Coverage Type", "What Was Covered"],
        [
            ["Functional", "Registration, valid login, product search, product details, add to cart, cart update/removal, checkout, order confirmation, logout."],
            ["UI validation", "Home, login, register, product list, product details, cart, checkout, and confirmation page controls."],
            ["Edge", "Invalid email, invalid password, existing email, mismatched password, empty search, missing billing fields, optional DOB fields."],
            ["Non-functional", "Product search response-time validation, reliability via waits, maintainability via POM, observability through logs/reports/screenshots."],
        ],
        "D9EAF7",
        8,
    )
    add_table(
        doc,
        ["Page", "Key UI Elements Validated"],
        [
            ["Home", "Search box, search button, login link, register link, cart link."],
            ["Login", "Email field, password field, Remember Me checkbox, login button, error message."],
            ["Register", "Gender radio buttons, first name, last name, email, password, confirm password, register button."],
            ["Product results", "Product list, product names, prices, product links."],
            ["Product details", "Title, price, quantity field when available, Add to Cart button."],
            ["Shopping cart", "Cart rows, product names, quantity input, terms checkbox, checkout button."],
            ["Checkout", "Billing fields, shipping options, payment options, Next buttons, Confirm button."],
            ["Confirmation", "Thank you heading, order details, continue button."],
        ],
        "FCE4D6",
        8,
    )

    add_heading(doc, "Bugs And Observations Logged", 1)
    add_para(
        doc,
        "These are application observations found during automation. The application under test is external, so they were not fixed in application code. They were handled in the automation framework and can still be logged as defects or improvement points.",
    )
    add_table(
        doc,
        ["ID", "Bug / Observation", "Steps To Reproduce", "Actual Result", "Expected Result", "Severity", "Automation Handling"],
        [
            ["BUG-001", "Empty search uses browser alert", "Open home page; keep search field empty; click Search.", "Browser alert appears: Please enter some search keyword.", "Inline validation should appear on the page.", "Low/Medium", "Alert is accepted and validated."],
            ["BUG-002", "Checkout sections change dynamically", "Add product to cart; start checkout; click Next through checkout steps.", "Buttons and sections appear after each step and older locators may fail.", "Stable IDs or automation-friendly hooks should be available.", "Medium", "Step-specific waits and fallback locators added."],
            ["BUG-003", "Product/cart behavior varies by product", "Search different products; add to cart; inspect quantity and cart rows.", "Quantity and cart row behavior can vary by product.", "Purchasable products should behave consistently.", "Medium", "Flexible cart validation and optional quantity handling added."],
            ["BUG-004", "Optional DOB fields are not present", "Open registration page; inspect available form fields.", "DOB fields expected by older tests are not present.", "Tests should not depend on optional/missing fields.", "Low", "DOB selection skipped when fields are absent."],
        ],
        "E4DFEC",
        7,
    )

    add_heading(doc, "Code Optimization And Efficiency Improvements", 1)
    add_bullets(
        doc,
        [
            "Smart UI handling: fallback locators, explicit waits, visibility checks, clickability checks, JavaScript click fallback, and scroll before interaction.",
            "Reusable flow utility: register unique customer, login, add product to cart, and complete checkout with less duplicated setup code.",
            "Data independence: unique test users reduce dependency on fixed credentials and stale environment state.",
            "Dynamic checkout handling: step-specific methods for billing, shipping, payment, payment information, and confirmation.",
            "Evidence and debugging: Log4j logs, Extent Reports, screenshots on failure, and TestNG/Surefire reports.",
            "Execution efficiency: critical suite runs first, user-story suite validates acceptance coverage, and full regression runs before release.",
        ],
    )

    add_heading(doc, "Optimized Regression Suite", 1)
    add_diagram(doc, "risk_prioritization.png", "Figure 4: Risk-based prioritization used to decide execution order.", 7.8)
    add_table(
        doc,
        ["Suite", "Purpose", "Recommended Use"],
        [
            ["testng-critical.xml", "Runs the highest-risk order placement path.", "Smoke testing and build validation."],
            ["testng-user-story.xml", "Runs functional, UI, edge, and performance user-story tests.", "User-story acceptance validation."],
            ["testng.xml", "Runs complete 31-test regression suite.", "Full release regression."],
            ["testng-parallel.xml", "Parallel execution option.", "Faster CI execution after core suite is stable."],
        ],
        "D9EAD3",
        8,
    )

    add_heading(doc, "Manual Testing Vs AI-Assisted Testing", 1)
    add_para(
        doc,
        "This comparison directly answers the evaluation point. AI helped speed up scenario discovery and regression planning, while the tester still validated business relevance and execution results.",
    )
    add_table(
        doc,
        ["Area", "Traditional Manual Approach", "AI-Assisted Approach"],
        [
            ["Requirement analysis", "Tester manually reads the user story and identifies scenarios.", "AI quickly breaks the story into functional, UI, edge, and non-functional scenarios."],
            ["Test case design", "Time-consuming and dependent on tester experience.", "Faster generation of coverage matrix, priorities, and expected results."],
            ["Edge cases", "May be missed under time pressure.", "AI suggests negative and edge cases such as empty search, invalid login, and missing billing."],
            ["Regression selection", "Often all tests are treated equally.", "AI helps separate critical, high, medium, and low priority."],
            ["Maintenance", "Locators and flows are fixed manually after failures.", "AI suggests fallback locators, reusable flows, and smart waits."],
            ["Documentation", "Created after implementation or sometimes skipped.", "AI helps generate structured documentation and traceability."],
            ["Execution strategy", "Manual sequencing by tester.", "AI helps design smoke, user-story, and full regression suites."],
            ["Risk assessment", "Based on tester judgment only.", "AI assists with risk and business impact mapping."],
        ],
        "FFF2CC",
        8,
    )

    add_heading(doc, "Execution Time Comparison", 1)
    add_table(
        doc,
        ["Activity / Metric", "Without AI / Traditional Effort", "With AI-Assisted Automation"],
        [
            ["Scenario identification", "4 to 6 hours estimated.", "Faster initial coverage design from the user story."],
            ["Test case writing", "1 to 2 days estimated.", "Structured scenario and test case matrix generated quickly, then human-reviewed."],
            ["Manual execution of 31 tests", "3 to 5 hours per cycle estimated.", "Full suite executed automatically in 681.1 seconds."],
            ["Critical validation", "Manual smoke execution still needs repeated effort.", "Critical suite gives faster feedback around the main order placement path."],
            ["Earlier full run", "31 tests with 13 failures due to brittle design.", "After refactor: 31 tests, 0 failures."],
            ["Main improvement", "High repeated debugging and maintenance effort.", "Stable, repeatable execution with logs, reports, screenshots, and data independence."],
        ],
        "FCE4D6",
        8,
    )

    add_heading(doc, "Risk-Based Test Prioritization", 1)
    add_table(
        doc,
        ["Priority", "Meaning", "Examples"],
        [
            ["Critical", "Directly impacts revenue or order placement.", "Login, product add to cart, checkout, confirmation."],
            ["High", "Important validation around critical flow.", "Product display, billing validation, shipping method, payment method."],
            ["Medium", "Important regression but not always blocking.", "Cart update, remove item, empty search, multiple products."],
            ["Low", "Lower business impact or optional behavior.", "Optional DOB fields and secondary UI behavior."],
        ],
        "E4DFEC",
        8,
    )
    add_numbers(
        doc,
        [
            "Run the critical order placement suite first.",
            "Run the user-story suite for UI, edge, and non-functional validation.",
            "Run the full regression suite before release.",
            "Run parallel or browser matrix only after the core suite is stable.",
        ],
    )

    add_heading(doc, "Observations On Coverage And Efficiency", 1)
    add_bullets(
        doc,
        [
            "The user story is covered end to end from login to order confirmation.",
            "Functional and UI validation are included at each major page.",
            "Negative and edge tests are included for login, registration, search, and checkout.",
            "Non-functional validation is included through product search response timing.",
            "Data-driven checkout and user data are supported through Excel.",
            "Critical-first execution reduces feedback time compared with running full regression every time.",
            "Smart UI handling reduces maintenance when the website UI changes frequently.",
        ],
    )

    add_heading(doc, "Critical Evaluation Of AI Effectiveness", 1)
    add_table(
        doc,
        ["Evaluation Area", "Assessment"],
        [
            ["Strengths", "AI accelerated scenario generation, risk analysis, regression optimization, locator resilience ideas, and documentation."],
            ["Limitations", "AI still required live execution feedback, real browser validation, and human decisions about bugs versus automation issues."],
            ["Human tester role", "Review generated scenarios, validate business relevance, approve risk prioritization, run tests, and review evidence."],
            ["Overall result", "AI was highly effective as an accelerator, but it did not replace QA judgment."],
        ],
        "D9EAF7",
        8,
    )

    add_heading(doc, "Final Execution Evidence", 1)
    add_table(
        doc,
        ["Artifact", "Location / Result"],
        [
            ["Final command", 'mvn -q test "-Dheadless=true"'],
            ["Final result", "Tests run: 31, Failures: 0, Errors: 0, Skipped: 0"],
            ["Execution time", "681.1 seconds"],
            ["Extent Report", "reports/ExtentReport_2026-05-27_19-19-09.html"],
            ["TestNG Report", "target/surefire-reports/index.html"],
            ["Test Summary", "target/surefire-reports/TestSuite.txt"],
            ["Logs", "logs/automation.log"],
            ["Screenshots", "screenshots/"],
            ["Excel Test Data", "test-data/TestData.xlsx"],
        ],
        "D9EAD3",
        8,
    )

    add_heading(doc, "Conclusion", 1)
    add_para(
        doc,
        "The project is complete and working according to the user story. It validates login, product discovery, add to cart, checkout, and order confirmation. It also includes functional, UI, edge, non-functional, data-driven, logging, reporting, screenshot, and optimized regression coverage.",
    )
    add_para(
        doc,
        "The final framework is stable, maintainable, and suitable for risk-based regression execution. The final verified result is 31 tests passed, 0 failed, 0 errors, and 0 skipped.",
    )

    doc.save(OUTPUT)
    print(OUTPUT)
    print(OUTPUT.stat().st_size)


if __name__ == "__main__":
    build()
